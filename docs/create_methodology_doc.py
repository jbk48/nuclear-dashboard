"""
create_methodology_doc.py
Generates Nuclear_Capacity_Dashboard_Methodology.docx using python-docx.
Run: python3 docs/create_methodology_doc.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "docs/Nuclear_Capacity_Dashboard_Methodology.docx"

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x37, 0x5E)   # primary header colour
BLUE      = RGBColor(0x2E, 0x75, 0xB6)   # accent blue
TEAL      = RGBColor(0x00, 0x70, 0x80)   # teal accent
GOLD      = RGBColor(0xC9, 0x9A, 0x00)   # gold / warning
RED       = RGBColor(0xC0, 0x39, 0x2B)   # red accent
GREEN     = RGBColor(0x1E, 0x8B, 0x4C)   # green accent
LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xF8)   # light blue-grey background
CALLOUT_BG= RGBColor(0xE8, 0xF4, 0xFD)   # callout box background
WARN_BG   = RGBColor(0xFF, 0xF3, 0xCD)   # warning/note background
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)
MID_GREY  = RGBColor(0x70, 0x70, 0x70)
TBL_HDR   = RGBColor(0x1A, 0x37, 0x5E)   # table header bg (navy)
TBL_ALT   = RGBColor(0xF2, 0xF7, 0xFD)   # table alternate row
DECLINE_C = RGBColor(0xC0, 0x39, 0x2B)
CONS_C    = RGBColor(0xE6, 0x7E, 0x22)
BASE_C    = RGBColor(0x27, 0x6F, 0xBF)
OPT_C     = RGBColor(0x1E, 0x8B, 0x4C)


# ── Helpers ───────────────────────────────────────────────────────────────────

def hex_to_str(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_to_str(color))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders: top, bottom, left, right each with color, sz, val."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{"insideH" if edge=="insideH" else edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'),    str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[edge].get('color', 'AAAAAA'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=None, color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if underline:
        run.underline = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing      = Pt(line)


def heading(doc, text, level=1, color=NAVY, size=None):
    """Add a styled heading paragraph."""
    sizes = {1: 20, 2: 15, 3: 12}
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(size or sizes.get(level, 12))
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after  = Pt(6)
        # Underline via border
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), hex_to_str(BLUE))
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)
    elif level == 2:
        pf.space_before = Pt(14)
        pf.space_after  = Pt(4)
    else:
        pf.space_before = Pt(10)
        pf.space_after  = Pt(3)
    return p


def body_para(doc, text='', before=4, after=4):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    if text:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
    para_spacing(p, before=before, after=after)
    return p


def bullet_para(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
    return p


def add_callout(doc, title, text, bg=CALLOUT_BG, title_color=BLUE):
    """Add a shaded callout box using a 1-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6.5)
    # Left border accent
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        tag.set(qn('w:sz'), '0')
        tag.set(qn('w:color'), 'auto')
        tcBorders.append(tag)
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '18')
    left.set(qn('w:color'), hex_to_str(title_color))
    tcBorders.append(left)
    tcPr.append(tcBorders)
    # Cell margins
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top','80'),('left','160'),('bottom','80'),('right','160')]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), val)
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)
    # Title
    p = cell.paragraphs[0]
    if title:
        run = p.add_run(title + '  ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = title_color
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK_GREY
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()  # spacer


def make_table(doc, headers, rows, col_widths=None, header_bg=TBL_HDR,
               stripe=True, font_size=9):
    """Create a styled table."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hr = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hr.cells, headers)):
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(font_size)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if col_widths and i < len(col_widths):
            cell.width = Inches(col_widths[i])

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        bg = TBL_ALT if (stripe and ri % 2 == 0) else WHITE
        for ci, (cell, val) in enumerate(zip(row.cells, row_data)):
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            if isinstance(val, tuple):  # (text, bold, color)
                run = p.add_run(val[0])
                run.bold = val[1] if len(val) > 1 else False
                run.font.color.rgb = val[2] if len(val) > 2 else DARK_GREY
            else:
                run = p.add_run(str(val))
                run.font.color.rgb = DARK_GREY
            run.font.size = Pt(font_size)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if col_widths and ci < len(col_widths):
                cell.width = Inches(col_widths[ci])

    doc.add_paragraph()  # spacer after table
    return table


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_enum_break())
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def docx_enum_break():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


def add_horizontal_rule(doc, color=BLUE):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_to_str(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def scenario_header_box(doc, label, color):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()


# ── Document builder ──────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Page setup: US Letter, 1-inch margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default paragraph style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK_GREY
    normal.paragraph_format.space_after = Pt(6)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # Top colour bar (fake with shaded table)
    cover_bar = doc.add_table(rows=1, cols=1)
    cover_bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    cb = cover_bar.cell(0, 0)
    set_cell_bg(cb, NAVY)
    p = cb.paragraphs[0]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    doc.add_paragraph()

    # Title block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_title.add_run('Nuclear Capacity Dashboard')
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY
    para_spacing(p_title, before=24, after=4)

    p_sub = doc.add_paragraph()
    r2 = p_sub.add_run('Methodology & User Guide')
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = BLUE
    para_spacing(p_sub, before=0, after=12)

    add_horizontal_rule(doc, BLUE)
    doc.add_paragraph()

    # Meta info table
    meta = doc.add_table(rows=4, cols=2)
    meta_data = [
        ('Version', '1.0  |  April 2026'),
        ('Baseline Year', '2024 (last full IAEA PRIS vintage)'),
        ('Projection Horizon', '2024 – 2050'),
        ('Primary Data Sources', 'IAEA PRIS · IAEA RDS-1/2 · IEA WEO 2024 · WNA Reactor Database'),
    ]
    for i, (label, val) in enumerate(meta_data):
        lc, vc = meta.rows[i].cells
        set_cell_bg(lc, LIGHT_BG)
        p1 = lc.paragraphs[0]
        r1 = p1.add_run(label)
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
        p2 = vc.paragraphs[0]
        r2 = p2.add_run(val)
        r2.font.size = Pt(10); r2.font.color.rgb = DARK_GREY

    doc.add_paragraph()
    add_horizontal_rule(doc, LIGHT_BG)

    # Scenario colour legend on cover
    p_leg = doc.add_paragraph()
    r_l = p_leg.add_run('Scenarios covered: ')
    r_l.bold = True; r_l.font.size = Pt(10); r_l.font.color.rgb = DARK_GREY
    for label, col in [('Decline', DECLINE_C), ('Conservative', CONS_C),
                       ('Base', BASE_C), ('Optimistic', OPT_C)]:
        r = p_leg.add_run(f'  ■ {label}')
        r.font.size = Pt(10); r.font.color.rgb = col; r.bold = True

    page_break(doc)

    # ── EXECUTIVE SUMMARY (PAGE 2) ────────────────────────────────────────────
    heading(doc, 'Executive Summary', level=1)
    body_para(doc, (
        'The Nuclear Capacity Dashboard is an interactive, scenario-based projection tool that '
        'tracks and forecasts global nuclear electricity generating capacity from 2024 to 2050. '
        'It is designed for energy analysts, policy professionals, and senior nuclear industry '
        'audiences who require a rigorous, transparent, and fully auditable view of how nuclear '
        'capacity may evolve under different policy and market conditions.'
    ))

    heading(doc, 'What the dashboard does', level=2)
    body_para(doc, (
        'This dashboard provides a bottom-up model of global nuclear capacity from 2024 through 2040, '
        'with a top-down projection of 2040–2050 installed capacity. Projections through 2040 are fully '
        'customizable based on retirement policy changes, assumptions about the likelihood of reactors in '
        'the nuclear pipeline being constructed, construction delays and more. In addition, the model can '
        'be customized down to the reactor level with direct overrides on reactor retirement and '
        'construction date, restarts, and uprates/downrates. These customizations can be made manually or '
        'using plain-English requests to an embedded Claude agent, which translates user hypotheticals '
        'into modelable assumptions.'
    ))
    body_para(doc, 'This means that the dashboard:')
    bullet_para(doc, 'Tracks 400+ individual reactor units by country, status, age, and capacity')
    bullet_para(doc, 'Projects retirements unit-by-unit using regulator-specific license renewal rules for 30+ countries')
    bullet_para(doc, 'Applies pipeline realization rates to Under Construction, Planned, and Proposed reactors')
    bullet_para(doc, 'Bridges to a top-down post-2040 growth assumption anchored to leading external benchmarks')
    bullet_para(doc, 'Provides four pre-defined scenarios plus full user customization via the Scenario Lab')

    heading(doc, 'The four scenarios at a glance', level=2)
    body_para(doc, (
        'The model includes four default scenarios that consider a range of assumptions about the future '
        'of nuclear power, based on life extension policies and pipeline realization. Users can select a '
        'given scenario and then layer on additional customizations if they so choose.'
    ))
    scen_rows = [
        ('Decline',       'AcceleratedRetirement', 'Low (UC only)',            '~10 GW/yr',  '~250 GW', 'IEA Low Nuclear Case',   "IEA's 'severe headwinds' case — a credible lower bound for a world where policy support fades."),
        ('Conservative',  'CurrentPolicy (~50%)',  'Medium (UC + Planned)',    '~35 GW/yr',  '~561 GW', 'IAEA Low Case (RDS-1)',  "IAEA's cautious-but-realistic view; widely cited by policymakers as a plausible floor."),
        ('Base',          'CurrentPolicy (~50%)',  'Medium (UC + Planned)',    '~30 GW/yr',  '~647 GW', 'IEA STEPS (WEO 2024)',   "IEA's central case conditioned on enacted policies — the most widely cited single reference."),
        ('Optimistic',    'ExtendedOperations',    'High (UC + Planned + Proposed)', '~57 GW/yr', '~992 GW', 'IAEA High Case (RDS-1)', "IAEA's high-ambition case; consistent with stated government expansion targets if fully delivered."),
    ]
    make_table(doc,
        ['Scenario', 'Life Extension Policy', 'Pipeline Realization', 'Post-2040 New Build', '2050 Capacity', 'Benchmark', 'Why this benchmark?'],
        scen_rows,
        col_widths=[0.7, 1.0, 1.0, 0.85, 0.7, 0.95, 1.3],
        font_size=8,
    )

    heading(doc, 'Key modeling principles', level=2)
    bullet_para(doc, 'Bottom-up through 2040: every retirement and new-build tracked unit by unit')
    bullet_para(doc, 'Top-down post-2040: builds a bridge between the bottom-up 2040 modeling and a top-down 2050 global capacity figure from either the IEA or IAEA, resulting in a single GW/year rate calibrated to the benchmark, then distributed regionally using IAEA/IEA proportions')
    bullet_para(doc, 'US in Base/Conservative: all US plants receive ExtendedOperations treatment, consistent with NRC license renewal data')
    bullet_para(doc, 'Binary pipeline realization: each reactor either fully enters service or is fully excluded (no partial capacity)')
    bullet_para(doc, '2024 baseline: last full IAEA PRIS vintage with known fleet composition and historical capacity record')

    page_break(doc)

    # ── TABLE OF CONTENTS (manual) ────────────────────────────────────────────
    heading(doc, 'Contents', level=1)
    toc_items = [
        ('1', 'Introduction & Purpose', '4'),
        ('2', 'Model Architecture', '5'),
        ('3', 'Data Sources, Baseline Year & Update Cadence', '7'),
        ('4', 'Retirement Framework', '10'),
        ('5', 'Pipeline Realization Framework', '13'),
        ('6', 'Scenarios in Detail', '15'),
        ('7', 'Benchmark Alignment', '23'),
        ('8', 'SMR Assumptions', '25'),
        ('9', 'Limitations & Caveats', '26'),
        ('10', 'Dashboard Tab Guide', '28'),
        ('', 'Appendix A: Glossary', '35'),
    ]
    for num, title, pg in toc_items:
        p = doc.add_paragraph()
        if num:
            r = p.add_run(f'{num}.  {title}')
        else:
            r = p.add_run(f'      {title}')
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY if num else DARK_GREY
        r.bold = bool(num)
        # dots + page number
        r2 = p.add_run(f'  {"." * max(1, 55 - len(num + title))}  {pg}')
        r2.font.size = Pt(9)
        r2.font.color.rgb = MID_GREY
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    page_break(doc)

    # ── SECTION 1: INTRODUCTION ───────────────────────────────────────────────
    heading(doc, '1.  Introduction & Purpose', level=1)
    body_para(doc, (
        'Nuclear power contributes approximately 10% of global electricity generation and remains '
        'a critical low-carbon baseload technology. Yet its long-term trajectory depends on decisions '
        'that vary enormously by jurisdiction: license renewals, construction pipelines, phase-out '
        'policies, and post-2040 investment commitments. No single public tool integrates all of '
        'these in a transparent, user-adjustable format.'
    ))
    body_para(doc, (
        'The Nuclear Capacity Dashboard fills this gap. It is built around a hybrid bottom-up / '
        'top-down projection engine, a 30-country retirement-rule database, and four independently '
        'calibrated scenarios that span the credible range of nuclear futures from deep decline '
        'to a sustained renaissance. Every assumption is documented here and, where possible, '
        'traceable to a named regulatory or institutional source.'
    ))

    heading(doc, '1.1  Intended Audience', level=2)
    bullet_para(doc, 'Energy and utility analysts tracking nuclear capacity for resource planning or investment thesis work')
    bullet_para(doc, 'Policy professionals assessing decarbonization pathways and technology mix scenarios')
    bullet_para(doc, 'Journalists and researchers requiring a transparent, citable quantitative framework')
    bullet_para(doc, 'Nuclear industry professionals who need a live tool for scenario sensitivity analysis')

    add_callout(doc,
        'Note on precision:',
        'Capacity figures are stated in gigawatts (GW) net. Individual reactor capacities are '
        'sourced from IAEA PRIS net electrical capacity ratings. Small rounding differences (±0.1 GW) '
        'may exist versus other published totals depending on the treatment of long-term shutdown units.',
        bg=CALLOUT_BG, title_color=BLUE)

    page_break(doc)

    # ── SECTION 2: MODEL ARCHITECTURE ────────────────────────────────────────
    heading(doc, '2.  Model Architecture', level=1)
    body_para(doc, (
        'The projection engine uses a two-phase hybrid approach that ensures unit-level fidelity '
        'in the near term while remaining tractable and benchmark-anchored in the long term.'
    ))

    heading(doc, '2.1  Historical Data Layer (2005–2024)', level=2)
    body_para(doc, (
        'The dashboard displays historical global and regional nuclear capacity sourced from IAEA PRIS '
        'for 2005–2024. This provides two decades of context for the projection, showing the post-2011 '
        'Fukushima dip, the gradual recovery of the fleet, and the current ~370–390 GW operating base. '
        'Historical data is stored in the local database and updated annually when IAEA PRIS publishes '
        'its revised figures.'
    ))

    heading(doc, '2.2  Bottom-Up Phase (2024–2040)', level=2)
    body_para(doc, (
        'Every reactor unit in the database contributes individually to the projection in this phase. '
        'Three drivers determine annual capacity:'
    ))

    # Architecture box
    arch_rows = [
        ('Retirements', 'Each operating reactor is assigned a retirement year from the retirement rule framework (Section 4). In the transition year, the unit\'s capacity is removed.'),
        ('New Builds (Pipeline)', 'Under Construction, Planned, and Proposed reactors enter service in their expected online year, weighted by their realization rate (Section 5). A construction delay adder can shift online years.'),
        ('Region Adders', 'Manual capacity adjustments (e.g., for specific projects not yet in PRIS) can be loaded per region per year. Used sparingly to correct known PRIS omissions.'),
    ]
    make_table(doc,
        ['Driver', 'Description'],
        arch_rows,
        col_widths=[1.5, 5.0],
        font_size=9,
    )

    body_para(doc, 'The annual formula per region is:')
    add_callout(doc, '',
        'capacity(year, region) = Σ[operating units not yet retired] + Σ[pipeline units × probability × (1 if online_year ≤ year)] + region_adders(year, region)',
        bg=LIGHT_BG, title_color=NAVY)

    heading(doc, '2.3  Top-Down Bridge (Post-2040)', level=2)
    body_para(doc, (
        'Bottom-up projections beyond 2040 become increasingly speculative: few reactors have '
        'construction start dates that far out, and even those that do face enormous policy uncertainty. '
        'The model therefore switches to a top-down growth parameterization:'
    ))
    body_para(doc, (
        'For each of the four scenarios, a 2050 calibration benchmark was selected from either the IEA '
        'or IAEA based on which benchmark appeared to best align with the kind of world the pre-2040 '
        'scenario indicated.'
    ))

    bridge_rows = [
        ('Decline',      '~10 GW/yr',  '~250 GW', 'IEA Low Nuclear Case (WEO 2024)'),
        ('Conservative', '~35 GW/yr',  '~561 GW', 'IAEA Low Case (RDS-1 2025)'),
        ('Base',         '~30 GW/yr',  '~647 GW', 'IEA STEPS (WEO 2024)'),
        ('Optimistic',   '~57 GW/yr',  '~992 GW', 'IAEA High Case (RDS-1 2025)'),
    ]
    make_table(doc,
        ['Scenario', 'Post-2040 Global New Build Rate', '2050 Target Capacity', 'Calibration Benchmark'],
        bridge_rows,
        col_widths=[1.3, 1.8, 1.5, 2.9],
        font_size=9,
    )

    body_para(doc, (
        'From there, we assume a consistent annual rate of new installed capacity between 2041 and 2050. '
        'A single post2040_global_growth_gw_per_year rate is set per scenario (calibrated to the '
        'benchmark — see Section 7). This global rate is distributed across regions using the '
        'proportions implied by IAEA/IEA regional forecasts. Unit-level retirements of the existing '
        '2024 fleet continue to be tracked individually through to 2050.'
    ))
    body_para(doc, (
        'The transition year is 2040 for all pre-defined scenarios. Continuity is guaranteed: '
        'capacity(2040) computed bottom-up equals capacity(2040) used as the starting point for '
        'top-down growth, by construction.'
    ))

    heading(doc, '2.4  What the Model Does Not Include', level=2)
    bullet_para(doc, 'Generation (TWh): only installed capacity (GW) is modelled')
    bullet_para(doc, 'Electricity market dynamics: no price signals, no dispatch modeling')
    bullet_para(doc, 'Carbon accounting: no direct CO₂ or lifecycle emissions')
    bullet_para(doc, 'Decommissioning timelines: the model tracks retirement year (plant stops generating) not decommissioning completion')
    bullet_para(doc, 'Probabilistic Monte Carlo: scenarios use fixed realization rates, not probability distributions')

    add_callout(doc,
        'Annual update cadence:',
        'The model is designed to be updated each year when IAEA PRIS releases its annual data '
        '(typically Q1). The update process involves: (1) refreshing the reactor table from IAEA PRIS, '
        '(2) re-running retirement-date derivation, (3) refreshing pipeline status, (4) re-calibrating '
        'post-2040 rates to the latest benchmark publications. A full update typically requires 1–2 days.',
        bg=WARN_BG, title_color=GOLD)

    page_break(doc)

    # ── SECTION 3: DATA SOURCES ───────────────────────────────────────────────
    heading(doc, '3.  Data Sources, Baseline Year & Update Cadence', level=1)

    heading(doc, '3.1  Why 2024 as the Baseline Year?', level=2)
    body_para(doc, (
        'IAEA PRIS releases a comprehensive reactor status update annually, with the most complete '
        'picture of the prior calendar year typically available by Q1 of the following year. '
        'The 2024 vintage (released in early 2025) is therefore the most recent year for which '
        'a reliable, globally consistent fleet snapshot exists. Using a mid-year snapshot would '
        'introduce ambiguity about which retirements and connections to include; end-of-year '
        'clarity makes the baseline reproducible and auditable.'
    ))
    body_para(doc, (
        'Historical capacity data from IAEA PRIS (2005–2024) is also displayed in the dashboard, '
        'providing two decades of context. The pre-2005 period is deliberately omitted: reactor '
        'records for early units are less complete and the fleet composition was materially '
        'different (especially the Soviet legacy fleet in Eastern Europe), reducing comparability.'
    ))

    heading(doc, '3.2  Primary Sources', level=2)
    body_para(doc, (
        'Sources are grouped below by role in the model. Within each group, Tier 1 sources '
        '(authoritative, primary) take precedence over Tier 2 (supplementary) where they conflict.'
    ))

    # Group 1
    p_grp = doc.add_paragraph()
    r_grp = p_grp.add_run('Group 1 — Fleet, Retirement & Historical Data')
    r_grp.bold = True; r_grp.font.size = Pt(10); r_grp.font.color.rgb = NAVY
    p_grp.paragraph_format.space_before = Pt(6); p_grp.paragraph_format.space_after = Pt(2)
    make_table(doc, ['Source', 'Full Name', 'What it provides', 'Update frequency', 'Tier'],
        [('IAEA PRIS',         'Power Reactor Information System',
          'Operating status, net capacity, commercial operation date, shutdown dates for every reactor worldwide',
          'Annual (Q1)', 'Tier 1'),
         ('NRC (US)',          'US Nuclear Regulatory Commission license renewal filings',
          'US-specific 20-year and subsequent license renewal (SLR) applications and approvals',
          'Continuous', 'Tier 1 (US)'),
         ('National regulators','Country-specific bodies (ASN, ONR, SSM, NSSC, etc.)',
          'Country extension policies, approval decisions, safety review outcomes',
          'As published', 'Tier 1 (country)'),
        ], col_widths=[1.0, 1.4, 2.3, 1.0, 0.8], font_size=8)

    # Group 2
    p_grp2 = doc.add_paragraph()
    r_grp2 = p_grp2.add_run('Group 2 — Pipeline Data')
    r_grp2.bold = True; r_grp2.font.size = Pt(10); r_grp2.font.color.rgb = NAVY
    p_grp2.paragraph_format.space_before = Pt(6); p_grp2.paragraph_format.space_after = Pt(2)
    make_table(doc, ['Source', 'Full Name', 'What it provides', 'Update frequency', 'Tier'],
        [('IAEA RDS-2 2024', 'Reference Data Series No. 2: Nuclear Power Plants: Status and Changes',
          'Pipeline reactor data (under construction, planned, proposed)',
          'Annual', 'Tier 1/2'),
         ('WNA Reactor DB',  'World Nuclear Association Reactor Database',
          'Pipeline reactor details, expected online years, reactor types',
          'Continuous', 'Tier 2'),
         ('Internet-scraped / press', 'Utility releases, parliamentary reports, news sources',
          'Pipeline projects not yet in PRIS/WNA; announced policy changes; restart timelines',
          'Ad hoc', 'Tier 2'),
        ], col_widths=[1.0, 1.4, 2.3, 1.0, 0.8], font_size=8)

    # Group 3
    p_grp3 = doc.add_paragraph()
    r_grp3 = p_grp3.add_run('Group 3 — Benchmarks & Scenarios')
    r_grp3.bold = True; r_grp3.font.size = Pt(10); r_grp3.font.color.rgb = NAVY
    p_grp3.paragraph_format.space_before = Pt(6); p_grp3.paragraph_format.space_after = Pt(2)
    make_table(doc, ['Source', 'Full Name', 'What it provides', 'Update frequency', 'Tier'],
        [('IAEA RDS-1 2025', 'Reference Data Series No. 1: Nuclear Power Reactors in the World',
          'Global and regional capacity benchmarks (Low/High cases) for 2030/2040/2050',
          'Annual', 'Benchmark'),
         ('IEA WEO 2024',    'World Energy Outlook 2024',
          'Global nuclear capacity benchmarks (STEPS, APS, Low Nuclear) for 2030/2040/2050',
          'Annual (Nov)', 'Benchmark'),
        ], col_widths=[1.0, 1.4, 2.3, 1.0, 0.8], font_size=8)

    heading(doc, '3.3  Tier 1 vs Tier 2 Retirement Dates', level=2)
    body_para(doc, (
        'Reactor retirement dates in the database are classified into two tiers:'
    ))
    add_callout(doc,
        'Tier 1 — Declared dates:',
        'The reactor has a publicly declared planned shutdown date (e.g., a regulatory decision, '
        'utility commitment, or legally binding phase-out mandate). This date is taken directly and '
        'used as-is. Examples: German phase-out units, UK Hinkley Point B, French regulatory decisions.',
        bg=CALLOUT_BG, title_color=BLUE)
    add_callout(doc,
        'Tier 2 — Modelled dates:',
        'No declared date exists. The retirement year is derived by applying the country\'s '
        'extension policy (baseline license years + allowable extension increments) to the '
        'commercial operation date. The derivation uses the retirement_rules table in the database, '
        'which is parameterized from official regulatory documentation for each country.',
        bg=CALLOUT_BG, title_color=TEAL)

    heading(doc, '3.4  Internet-Scraped and Supplementary Data', level=2)
    body_para(doc, (
        'For pipeline reactors — particularly in China, India, Russia, and emerging markets — '
        'the PRIS and WNA records may lag announced developments by 6–18 months. The database '
        'is therefore supplemented by:'
    ))
    bullet_para(doc, 'Utility investor relations releases (e.g., EDF, KEPCO, Rosatom)')
    bullet_para(doc, 'Parliamentary and government energy planning documents')
    bullet_para(doc, 'IAEA country mission reports (IRRS, INIR)')
    bullet_para(doc, 'Industry news wires (World Nuclear News, Nuclear Engineering International)')
    body_para(doc, (
        'Internet-scraped data is classified as Tier 2 and treated with higher uncertainty. '
        'Where a Tier 2 source conflicts with a Tier 1 source, the Tier 1 value prevails. '
        'All data sources are logged in the data_sources_log table in the database.'
    ))

    heading(doc, '3.5  Annual Update Process', level=2)
    update_rows = [
        ('Step 1', 'Q1', 'Download IAEA PRIS annual data extract. Identify retired, restarted, and newly connected units.'),
        ('Step 2', 'Q1', 'Re-run retirement date derivation (retirement.py). Cross-check Tier 1 declared dates against regulatory announcements.'),
        ('Step 3', 'Q1–Q2', 'Refresh pipeline (under construction, planned, proposed) from RDS-2, WNA, and internet sources.'),
        ('Step 4', 'Q2', 'Re-calibrate post-2040 benchmark rates against latest IEA WEO and IAEA RDS-1 publications.'),
        ('Step 5', 'Q2', 'Validate model outputs against IAEA/IEA published 2030 estimates (sanity check within ±5%).'),
        ('Step 6', 'Ongoing', 'Monitor for major policy changes (e.g., new phase-out mandates, new license renewals) and apply point overrides.'),
    ]
    make_table(doc,
        ['Step', 'Timing', 'Action'],
        update_rows,
        col_widths=[0.6, 0.6, 5.3],
        font_size=9,
    )

    page_break(doc)

    # ── SECTION 4: RETIREMENT FRAMEWORK ──────────────────────────────────────
    heading(doc, '4.  Retirement Framework', level=1)
    body_para(doc, (
        'The retirement framework is the most consequential single component of the model. '
        'For the current operating fleet, retirements account for the majority of near-term '
        'capacity loss. The framework uses three extension policies, applied globally or '
        'per-country, and is underpinned by a 30-country regulatory database.'
    ))

    heading(doc, '4.1  Three Extension Policies', level=2)

    pol_rows = [
        ('AcceleratedRetirement', 'DECLINE scenario',
         'No extensions beyond already-approved (Tier 1) dates. Tier 2 reactors retire at their baseline license expiry (i.e., zero extension assumed). Reflects a regulatory or political environment that does not permit or fund further life extension.',
         DECLINE_C),
        ('CurrentPolicy', 'BASE / CONSERVATIVE scenarios',
         'Historical extension rates apply. Roughly 50% of eligible reactors receive one standard extension increment (per country rules). This reflects observed behavior: not all eligible units apply for or receive renewals. For the US, all plants are treated as ExtendedOperations because NRC license renewal data shows virtually all operational plants have already obtained 60-year licenses with SLR applications in progress.',
         BASE_C),
        ('ExtendedOperations', 'OPTIMISTIC scenario',
         'All eligible reactors receive the maximum life extension permitted under current law (e.g., up to 80 years for the US, 60 years for most European units, 60 years for Russian VVERs). Reflects a policy environment that actively incentivises long-term operation — consistent with stated intentions of multiple governments.',
         OPT_C),
    ]
    for code, used_in, desc, col in pol_rows:
        scenario_header_box(doc, f'{code}  —  used in {used_in}', col)
        p = doc.add_paragraph()
        r = p.add_run(desc)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_GREY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(8)

    heading(doc, '4.2  US Base Scenario: ExtendedOperations Treatment', level=2)
    add_callout(doc,
        'Key transparency note:',
        'In both the Base and Conservative scenarios, the United States fleet is assigned '
        'ExtendedOperations even though the global extension policy is CurrentPolicy. This is because '
        'NRC license renewal data shows that virtually all currently operating US plants have already '
        'obtained 60-year licenses, with subsequent license renewal (SLR) applications filed for '
        '80-year operation. Applying a 50% "CurrentPolicy" haircut to the US fleet would understate '
        'actual regulatory approvals already granted. This treatment is documented transparently '
        'in the scenario_regional_params table (pipeline_rate_source field).',
        bg=WARN_BG, title_color=GOLD)

    heading(doc, '4.3  Country-Level Retirement Rules', level=2)
    body_para(doc, (
        'The retirement_rules table contains parameters for 30+ countries/regions. '
        'Key columns: baseline_license_years (initial operating license), typical_extension_years '
        '(standard renewal increment), max_possible_life_years (ceiling under current law), '
        'extension_unit_years (granularity of extension steps). Sources are regulatory documents '
        'from each national authority.'
    ))

    # Key country table (selected subset, not all 40+ rows including aliases)
    country_rows = [
        ('United States', '40', '20 (×2 possible)', '80', '20', 'NRC 10 CFR 54; SLR in progress for most plants'),
        ('France',        '40', '10 (periodic)',     '60', '10', 'ASN 10-year safety reviews; 60yr under exploration'),
        ('China',         '40', '0 (not routine)',   '50', '10', '40yr design life; extensions not yet standard; young fleet'),
        ('Russia',        '30', '15',                '60', '15', 'Rosatom 30yr initial + 15yr extensions; 60yr target'),
        ('Japan',         '40', '20',                '80', '20', 'NRA post-Fukushima 40yr default; 2023 law allows beyond 60yr'),
        ('South Korea',   '40', '20',                '60', '20', 'NSSC 40yr initial; case-by-case extensions'),
        ('United Kingdom','35', '10',                '55', '10', 'ONR lifecycle license; ~40-50yr typical'),
        ('Canada',        '30', '30',                '60', '10', '~30yr initial; CANDU refurbishment adds ~30yr'),
        ('Germany',       '—',  '0 (phased out)',    '—',  '—',  'Phase-out completed April 2023; no extensions'),
        ('Sweden',        '40', '20',                '60', '20', 'SSM; extensions planned for Ringhals/Forsmark'),
        ('Belgium',       '40', '10',                '50', '10', 'FANC; current policy 50yr for Doel/Tihange'),
        ('Finland',       '40', '20',                '60', '20', 'STUK; Loviisa extensions approved'),
        ('India',         '40', '20',                '60', '20', 'AERB 40yr; case-by-case extensions'),
        ('Ukraine',       '30', '20',                '60', '20', 'Soviet VVER; 30yr initial + extensions'),
        ('Taiwan',        '40', '0',                 '40', '0',  'Phase-out policy; no extensions approved'),
        ('UAE',           '60', '0',                 '60', '0',  'Barakah: 60yr design life; new fleet'),
        ('Hungary',       '30', '20',                '60', '20', 'Paks 1-4 extended to ~50yr; Paks II under construction'),
        ('Slovakia',      '30', '20',                '60', '20', 'UJD; Bohunice/Mochovce extensions'),
        ('Romania',       '30', '20',                '60', '20', 'Cernavoda refurbishment planned'),
        ('South Africa',  '40', '20',                '60', '20', 'NNR; Koeberg extended to 60yr'),
    ]
    make_table(doc,
        ['Country', 'Baseline Licence (yrs)', 'Extension (yrs)', 'Max Life (yrs)', 'Increment (yrs)', 'Key Notes'],
        country_rows,
        col_widths=[1.15, 0.85, 0.9, 0.75, 0.75, 3.1],
        font_size=8,
    )
    body_para(doc, '* Germany phase-out complete; units removed from active fleet.',
              before=0, after=8)

    page_break(doc)

    # ── SECTION 5: PIPELINE REALIZATION ──────────────────────────────────────
    heading(doc, '5.  Pipeline Realization Framework', level=1)
    body_para(doc, (
        'The pipeline realization framework determines what fraction of reactors that are not '
        'yet operational actually enter service by their expected online year. The framework '
        'distinguishes three pipeline status categories and applies binary realization rates '
        '(0% or 100% per reactor, not fractional).'
    ))

    heading(doc, '5.1  Pipeline Status Categories', level=2)
    status_rows = [
        ('Under Construction (UC)', 'Major civil construction underway; first concrete poured or equivalent milestone reached. Verified from IAEA PRIS construction start dates.', '95–100%', 'Very high — physical assets committed'),
        ('Planned',                 'Firm government or utility commitment: site licensed, significant engineering investment made, financial close expected. Typically includes projects with construction scheduled within 3–5 years.', '0–100%', 'Medium — depends on policy and financing'),
        ('Proposed',                'Early-stage: announced, on national energy plans, or feasibility study stage. Significant permitting, financing, and regulatory hurdles remain.', '0–100%', 'Low in most scenarios — high uncertainty'),
    ]
    make_table(doc,
        ['Status Category', 'Definition', 'Typical Rate Range', 'Rationale'],
        status_rows,
        col_widths=[1.4, 2.6, 1.2, 1.8],
        font_size=9,
    )

    heading(doc, '5.2  The Three Pipeline Presets', level=2)
    body_para(doc, (
        'Three named presets simplify the lever interface. The same preset definitions apply to '
        'both large reactors (>300 MW) and small modular reactors (SMRs):'
    ))
    preset_rows = [
        ('High',   '100%', '100%', '100%', 'All announced projects complete; used in Optimistic scenario. Aggressive but represents stated government targets (China, Russia, India, South Korea combined pipeline).'),
        ('Medium', '100%', '100%',   '0%', 'UC and firm Planned projects complete; Proposed excluded. Used in Base and Conservative scenarios. Reflects observed track record of projects with committed financing.'),
        ('Low',    '100%',   '0%',   '0%', 'Only units already under construction complete. Used in Decline scenario. Conservative: accounts for high cancellation rates of Planned projects in most Western markets historically.'),
    ]
    make_table(doc,
        ['Preset', 'UC Rate', 'Planned Rate', 'Proposed Rate', 'Rationale'],
        preset_rows,
        col_widths=[0.8, 0.7, 0.85, 0.85, 4.3],
        font_size=9,
    )

    heading(doc, '5.3  Why Binary (Not Probabilistic) Realization?', level=2)
    body_para(doc, (
        'The pipeline_probability field exists in the database and is displayed in the reactor '
        'detail tables. However, the four preset scenarios use binary realization rates (0 or 1) '
        'rather than fractional probabilities for two reasons:'
    ))
    bullet_para(doc, 'Scenario clarity: each scenario represents a coherent policy world, not a probability-weighted average. A "Base" scenario that shows 0.6× of every reactor is less interpretable than one where a defined set of reactors either does or does not build.')
    bullet_para(doc, 'Benchmark anchoring: the post-2040 bridge rates are calibrated against IAEA/IEA totals; using fractional rates in the bottom-up phase would require recalibration for every intermediate value.')
    body_para(doc, (
        'The Scenario Lab allows users to set continuous pipeline rates (e.g., 40% Proposed '
        'realization) for custom scenarios, where individual reactor probabilities are applied directly.'
    ))

    heading(doc, '5.4  Construction Delay Adder', level=2)
    body_para(doc, (
        'Each scenario includes a construction_delay_adder_years parameter that shifts all '
        'pipeline reactor online years forward by the specified amount. This captures systemic '
        'supply-chain and labour constraints:'
    ))
    delay_rows = [
        ('Decline',      '3 years', 'High delay environment — reflects recent Western experience (Hinkley Point C, Vogtle, Flamanville)'),
        ('Conservative', '2 years', 'Moderate delays — some improvement in supply chain but structural constraints persist'),
        ('Base',         '0 years', 'No additional delay — modelled expected online years taken as stated'),
        ('Optimistic',   '0 years', 'No additional delay — assumes policy and supply chain support for timely delivery'),
    ]
    make_table(doc,
        ['Scenario', 'Delay Adder', 'Rationale'],
        delay_rows,
        col_widths=[1.2, 1.1, 4.2],
        font_size=9,
    )

    page_break(doc)

    # ── SECTION 6: SCENARIOS IN DETAIL ────────────────────────────────────────
    heading(doc, '6.  Scenarios in Detail', level=1)
    body_para(doc, (
        'This section documents each scenario\'s full parameter set and rationale. '
        'All scenario parameters are stored in the scenarios and scenario_regional_params '
        'tables in the database and are directly used by the projection engine — '
        'there are no undocumented adjustments.'
    ))

    # ── DECLINE ──
    scenario_header_box(doc, '6.1  Decline Scenario', DECLINE_C)
    body_para(doc, (
        'The Decline scenario represents a world where policy support for nuclear fades, '
        'construction programs stall, and life extensions are not pursued. It anchors to '
        'the IEA Low Nuclear Case, which represents approximately 250 GW of global capacity '
        'by 2050 — a scenario IEA describes as where "nuclear power faces severe constraints." '
        'This is not a "nuclear disappears" scenario; it is a world of managed decline from '
        'the current ~380 GW base to a residual fleet concentrated in geographies with long-term '
        'state support (China, Russia, India).'
    ))
    bullet_para(doc, 'Life extension policy: ', bold_prefix='Life extension policy: ', )
    p = doc.paragraphs[-1]
    p.clear()
    run = p.add_run('Life extension policy: ')
    run.bold = True; run.font.size = Pt(10); run.font.color.rgb = DARK_GREY
    run2 = p.add_run('AcceleratedRetirement — No Tier 2 extensions. Every reactor retires at its initial license expiry or declared date, whichever comes first. Already-approved extensions (Tier 1) are honored.')
    run2.font.size = Pt(10); run2.font.color.rgb = DARK_GREY
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

    decline_detail = [
        ('Pipeline (large reactors)', 'Low: 100% UC / 0% Planned / 0% Proposed'),
        ('Pipeline (SMRs)', 'Low: 100% UC / 0% Planned / 0% Proposed'),
        ('Construction delay adder', '3 years added to all pipeline online years'),
        ('Post-2040 new build', '~10 GW/yr globally (barely offsetting retirements)'),
        ('US extension policy', 'AcceleratedRetirement (consistent with global policy)'),
        ('Post-2040 China new build', '~2.8 GW/yr (much lower than other scenarios)'),
        ('2050 target capacity', '~250 GW (IEA Low Nuclear Case)'),
    ]
    make_table(doc,
        ['Parameter', 'Value'],
        decline_detail,
        col_widths=[2.2, 4.3],
        font_size=9,
    )
    body_para(doc, (
        'Benchmark alignment: The IEA Low Nuclear Case (WEO 2024, ~250 GW by 2050) represents '
        '"a challenging environment in which nuclear faces significant headwinds." Key comparators: '
        'IAEA Low Case (2050) = 561 GW, indicating the IEA Low Nuclear is more pessimistic than '
        'even the IAEA conservative view. This scenario is included to bracket the lower end of '
        'credible outcomes, not as a base expectation.'
    ))
    add_callout(doc,
        'Why this matters:',
        'Even in Decline, China and Russia continue adding capacity under state-directed programs. '
        'The scenario does not assume zero new build globally — it assumes Western cancellations '
        'dominate, and Asian growth is insufficient to offset total retirements.',
        bg=RGBColor(0xFF, 0xEB, 0xEB), title_color=DECLINE_C)

    doc.add_paragraph()

    # ── CONSERVATIVE ──
    scenario_header_box(doc, '6.2  Conservative Scenario', CONS_C)
    body_para(doc, (
        'The Conservative scenario reflects cautious but not catastrophic nuclear futures: '
        'existing reactors operate to historically observed (not maximum) lifetimes, committed '
        'pipeline projects complete with moderate delays, and post-2040 new build grows steadily '
        'but primarily in Asia. This anchors to the IAEA Low Case (~561 GW by 2050).'
    ))
    cons_detail = [
        ('Life extension policy', 'CurrentPolicy: ~50% of eligible reactors receive one standard extension increment per country rules'),
        ('US exception', 'ExtendedOperations (all US plants treated as 60-80yr) — NRC data shows virtually all have already obtained 60yr licenses'),
        ('Pipeline (large reactors)', 'Medium: 100% UC / 100% Planned / 0% Proposed'),
        ('Pipeline (SMRs)', 'Medium: 100% UC / 100% Planned / 0% Proposed'),
        ('Construction delay adder', '2 years added to all pipeline online years'),
        ('Post-2040 new build', '~35 GW/yr globally; ~8.1 GW/yr China; ~4.9 GW/yr US; ~3.9 GW/yr South Asia'),
        ('2050 target capacity', '~561 GW (IAEA Low Case, RDS-1 2025)'),
    ]
    make_table(doc,
        ['Parameter', 'Value'],
        cons_detail,
        col_widths=[2.2, 4.3],
        font_size=9,
    )
    body_para(doc, (
        'Key country assumptions: France operates Tier 1 fleet (already at ~60yr policy), '
        'no new large reactor connections assumed until Flamanville-3 (already UC). '
        'Germany: phase-out complete, no reversal modelled. Japan: conservative restart rate — '
        'only units with existing or imminent approvals assumed to return. Eastern Europe: '
        'modest new build (Poland, Czech Republic planned units included at Planned rate).'
    ))
    add_callout(doc,
        'Why Conservative ≠ Decline:',
        'Conservative includes 100% Planned realization (unlike Decline\'s 0%), the 2-year '
        'delay adder (vs. 3 years), and higher post-2040 new build (~35 vs ~10 GW/yr). '
        'The cumulative effect over 25 years is ~310 GW more capacity by 2050.',
        bg=RGBColor(0xFF, 0xF3, 0xE0), title_color=CONS_C)

    doc.add_paragraph()

    # ── BASE ──
    scenario_header_box(doc, '6.3  Base Scenario (Central Case)', BASE_C)
    body_para(doc, (
        'The Base scenario represents the most likely trajectory under current stated government '
        'policies and observed industry behavior. It uses the same extension policy and pipeline '
        'preset as Conservative, but removes construction delays (no adder) and has a slightly '
        'higher post-2040 growth rate, anchoring to IEA STEPS (~647 GW by 2050).'
    ))
    body_para(doc, (
        'IEA STEPS (Stated Policies Scenario) is the IEA\'s central case for energy — it takes '
        'governments at their word on announced and legislated policies, without assuming additional '
        'ambition. For nuclear, STEPS implies steady growth in Asia, moderate fleet life in the US '
        'and France, and limited new build elsewhere.'
    ))
    base_detail = [
        ('Life extension policy', 'CurrentPolicy globally; ExtendedOperations for US'),
        ('Pipeline (large reactors)', 'Medium: 100% UC / 100% Planned / 0% Proposed'),
        ('Pipeline (SMRs)', 'Medium: 100% UC / 100% Planned / 0% Proposed'),
        ('Construction delay adder', '0 years (stated timelines taken as given)'),
        ('Post-2040 global new build', '~30 GW/yr globally (~28 GW calibrated value)'),
        ('Post-2040 China new build', '~8.0 GW/yr'),
        ('Post-2040 US new build', '~4.8 GW/yr'),
        ('Post-2040 South Asia', '~3.8 GW/yr (India-led)'),
        ('Post-2040 Russia', '~3.2 GW/yr'),
        ('Post-2040 East Asia', '~1.3 GW/yr (South Korea, Japan)'),
        ('2050 target capacity', '~647 GW (IEA STEPS, WEO 2024)'),
    ]
    make_table(doc,
        ['Parameter', 'Value'],
        base_detail,
        col_widths=[2.2, 4.3],
        font_size=9,
    )
    add_callout(doc,
        'Why STEPS and not IAEA Low?',
        'IAEA Low (561 GW) and IEA STEPS (647 GW) bracket a similar "cautious progress" space. '
        'We use STEPS as the Base because it is explicitly conditioned on current policies — '
        'the most interpretable central case. IAEA Low is assigned to Conservative because it '
        'better reflects partial extension and conservative pipeline realization.',
        bg=RGBColor(0xE8, 0xF2, 0xFF), title_color=BASE_C)

    doc.add_paragraph()

    # ── OPTIMISTIC ──
    scenario_header_box(doc, '6.4  Optimistic Scenario', OPT_C)
    body_para(doc, (
        'The Optimistic scenario models a sustained nuclear renaissance: maximum life extensions '
        'globally, full pipeline realization (including speculative proposed projects), no '
        'construction delays, and high post-2040 new build. It anchors to the IAEA High Case '
        '(~992 GW by 2050).'
    ))
    body_para(doc, (
        'This scenario is not implausible: multiple governments (US, France, UK, South Korea, '
        'Japan) have announced nuclear expansion targets in recent years. AI and data-center '
        'electricity demand growth has materially increased corporate PPA interest in nuclear. '
        'The scenario does, however, require simultaneously high political will, supply chain '
        'capacity, and financing — conditions that have historically not coexisted globally.'
    ))
    opt_detail = [
        ('Life extension policy', 'ExtendedOperations: all eligible reactors reach maximum permitted life (e.g., 80yr US, 60yr most of Europe, 60yr Russia)'),
        ('Pipeline (large reactors)', 'High: 100% UC / 100% Planned / 100% Proposed'),
        ('Pipeline (SMRs)', 'High: 100% UC / 100% Planned / 100% Proposed'),
        ('Construction delay adder', '0 years — optimistic delivery'),
        ('Post-2040 global new build', '~57 GW/yr globally (~54 GW calibrated value)'),
        ('Post-2040 China new build', '~15.2 GW/yr (>1 standard reactor per month)'),
        ('Post-2040 US new build', '~9.1 GW/yr'),
        ('Post-2040 South Asia', '~7.3 GW/yr'),
        ('Post-2040 Russia', '~6.1 GW/yr'),
        ('Post-2040 East Asia', '~2.4 GW/yr'),
        ('Post-2040 Western Europe (France)', '~1.8 GW/yr'),
        ('2050 target capacity', '~992 GW (IAEA High Case, RDS-1 2025)'),
    ]
    make_table(doc,
        ['Parameter', 'Value'],
        opt_detail,
        col_widths=[2.2, 4.3],
        font_size=9,
    )
    add_callout(doc,
        'Stress test:',
        '~57 GW/yr of new nuclear additions by the 2040s equates to roughly 45–55 new standard '
        'reactors per year globally (at ~1.0–1.2 GW each), or a much larger number of SMR units. '
        'For context, the peak historical build rate (1980s) was ~30–35 GW/yr globally. '
        'IAEA High requires surpassing that peak — ambitious but cited in IAEA publications '
        'as achievable under a sustained policy push.',
        bg=RGBColor(0xE8, 0xF5, 0xEE), title_color=OPT_C)

    page_break(doc)

    # ── SECTION 7: BENCHMARK ALIGNMENT ────────────────────────────────────────
    heading(doc, '7.  Benchmark Alignment', level=1)
    body_para(doc, (
        'The four scenarios are calibrated against published IAEA and IEA benchmarks. '
        'The table below shows how model outputs compare to these benchmarks at key milestones. '
        'Small deviations from the benchmark at 2050 are expected because the model\'s bottom-up '
        'phase (2024–2040) is fixed by fleet data; only the post-2040 rate is calibrated.'
    ))

    bench_rows = [
        ('IEA Low Nuclear Case', '—',    '~250 GW',  'Decline',      '~250 GW',  'IEA WEO 2024'),
        ('IAEA Low Case',        '425 GW', '519 GW', 'Conservative', '~561 GW',  'RDS-1 2025'),
        ('IEA STEPS',            '513 GW', '586 GW', 'Base',         '~647 GW',  'IEA WEO 2024'),
        ('IEA APS',              '551 GW', '748 GW', '(Reference)',   '~874 GW', 'IEA WEO 2024'),
        ('IAEA High Case',       '445 GW', '710 GW', 'Optimistic',   '~992 GW',  'RDS-1 2025'),
    ]
    make_table(doc,
        ['Benchmark', '2030 (GW)', '2040 (GW)', 'Matched Scenario', '2050 Target (GW)', 'Source'],
        bench_rows,
        col_widths=[1.6, 0.8, 0.8, 1.3, 1.3, 1.2],
        font_size=9,
    )

    body_para(doc, (
        'Note on IEA APS: The Announced Pledges Scenario is shown as a reference point '
        'but is not directly matched by a scenario. Its 2050 value (~874 GW) falls '
        'between the Base and Optimistic scenarios. Users can approximate this level '
        'using the Scenario Lab with high pipeline realization and partial Proposed inclusion.'
    ))

    heading(doc, '7.1  Convergence Method', level=2)
    body_para(doc, (
        'Each scenario\'s post-2040 global_growth_gw_per_year is calibrated by solving for the '
        'value that hits the benchmark\'s 2050 total, given the model\'s projected 2040 capacity. '
        'Because the 2040 capacity changes with fleet data updates, these rates are re-calibrated '
        'annually. The rates stored in the database are the calibrated values as of the 2024 data vintage.'
    ))
    body_para(doc, (
        'Regional allocation uses the proportions from IAEA/IEA regional breakdowns '
        '(Central & Eastern Asia, Eastern Europe, Northern America, etc.) re-mapped to the '
        'dashboard\'s 12 regions. The mapping is documented in the scenario_regional_params table.'
    ))

    add_callout(doc,
        'Important caveat:',
        'The 2030 and 2040 column values in the benchmark table above are from the external '
        'benchmark publications. The model\'s own 2030 and 2040 outputs are determined by the '
        'bottom-up phase and may differ from the benchmark\'s intermediate milestones, particularly '
        'if the benchmark uses different fleet assumptions for retirements or pipeline. '
        'Only the 2050 terminal value is explicitly calibrated.',
        bg=WARN_BG, title_color=GOLD)

    page_break(doc)

    # ── SECTION 8: SMR ────────────────────────────────────────────────────────
    heading(doc, '8.  Small Modular Reactor (SMR) Assumptions', level=1)
    body_para(doc, (
        'SMRs (reactors with net capacity ≤300 MW per unit) are tracked separately in the database '
        'using the is_smr flag. As of the 2024 baseline, very few SMRs are operational globally '
        '(a handful of demonstration units in China and Russia). The pipeline includes a modest '
        'number of planned and proposed projects.'
    ))

    heading(doc, '8.1  Current Treatment', level=2)
    bullet_para(doc, 'SMRs in Under Construction or Planned status: treated identically to large reactors — they enter service at their expected online year, subject to the SMR realization preset')
    bullet_para(doc, 'SMR pipeline preset: same H/M/L structure as large reactors (100/100/100%, 100/100/0%, 100/0/0%)')
    bullet_para(doc, 'Post-2040 SMR contribution: currently folded into the global post-2040 new build rate (not split out separately in the pre-defined scenarios)')

    heading(doc, '8.2  Limitations', level=2)
    body_para(doc, (
        'The SMR landscape is evolving rapidly (NuScale, GE-Hitachi BWRX-300, Rolls-Royce SMR, '
        'TerraPower Natrium, etc.). The dashboard does not model commercial SMR deployment curves '
        'separately from large reactor pipelines, for two reasons:'
    ))
    bullet_para(doc, 'Very few SMR designs have reached commercial operation; deployment timelines are highly speculative beyond 2030')
    bullet_para(doc, 'The post-2040 top-down rate implicitly includes SMR deployment — it is calibrated to the same benchmark that includes IAEA/IEA SMR projections')
    body_para(doc, (
        'A future update may add an explicit SMR acceleration module (accelerated_gw_per_year from '
        'a specified year, modifiable via the Scenario Lab). This is already partially implemented '
        'in the levers interface.'
    ))

    add_callout(doc,
        'SMR note for advanced users:',
        'In the Scenario Lab, you can set a separate SMR pipeline preset (e.g., High for SMRs '
        'even if Large is Medium). This captures the view that near-term SMR commitments may '
        'have higher execution risk than currently operating large reactor construction projects.',
        bg=CALLOUT_BG, title_color=TEAL)

    page_break(doc)

    # ── SECTION 9: LIMITATIONS ────────────────────────────────────────────────
    heading(doc, '9.  Limitations & Caveats', level=1)
    body_para(doc, (
        'No projection tool eliminates uncertainty. The following limitations are inherent '
        'in the methodology and should be understood by users before drawing policy conclusions.'
    ))

    # Ordered High → Medium → Low; binary realization removed (it is a design choice, documented in §5, not a limitation)
    lim_rows = [
        ('Post-2040 top-down',         'Post-2040 growth is a single global rate calibrated to a benchmark. It cannot anticipate regional surprises (e.g., sudden German reversal, unexpected Chinese slowdown).', 'High',              'Compare multiple scenarios; treat post-2040 as indicative order-of-magnitude'),
        ('Retirement date accuracy',   'Tier 2 retirement dates are derived from regulatory rules that may change. New life extension policies, phase-out decisions, or ownership changes can shift retirements by 10–20 years.', 'High',  'Use What-If retirement overrides in the Scenario Lab'),
        ('SMR uncertainty',            'SMR deployment timelines are highly speculative. The post-2040 rate includes SMR contribution but cannot disaggregate it reliably.', 'High (post-2030)',                                          'Treat post-2040 optimistic cases with appropriate skepticism for SMR-dependent trajectories'),
        ('Pipeline data lag',          'PRIS and WNA can lag real-world announcements by 6–18 months, particularly in China and Russia.',                                        'Medium',                                                  'Internet-scraped sources supplement but do not eliminate this lag'),
        ('Annual vintage',             'Data reflects end-2024 IAEA PRIS. Major 2025+ events (new approvals, cancellations, restarts) are not in the model unless manually entered via what-if overrides.', 'Medium',                    'Use the What-If Builder to add known 2025+ developments'),
        ('No market dynamics',         'The model does not respond to electricity prices, carbon prices, or gas competition. Retirements are rule-based, not economic.',          'Low–Medium',                                              'Externally assess economic retirement risk for price-exposed markets (Germany 2011–2023 as a cautionary example)'),
        ('Capacity factor / generation','Only installed capacity (GW) is modeled. Generation (TWh) depends on capacity factors, which vary by country and over reactor lifetime.', 'Low (for capacity questions)',                          'Multiply by national average capacity factors (~85–92% for baseload nuclear) to estimate generation'),
    ]
    make_table(doc,
        ['Limitation', 'Description', 'Materiality', 'Mitigation'],
        lim_rows,
        col_widths=[1.3, 2.4, 0.9, 1.9],
        font_size=8,
    )

    page_break(doc)

    # ── SECTION 10: TAB GUIDE ─────────────────────────────────────────────────
    heading(doc, '10.  Dashboard Tab Guide', level=1)
    body_para(doc, (
        'The dashboard has seven tabs, each designed for a specific analytical question. '
        'This section explains what each tab shows, how to use its controls, and '
        'what conclusions it supports.'
    ))

    tabs = [
        {
            'num': '10.1', 'name': 'Global Projection',
            'what': (
                'Displays total global nuclear capacity (GW) from 2005 to 2050, overlaying historical '
                'data (solid line) with the projected trajectory (dashed line from 2024). Up to four '
                'external benchmark lines (IAEA Low/High, IEA STEPS, IEA APS, IEA Low Nuclear) can be '
                'toggled on or off. A second scenario can be overlaid for direct comparison.'
            ),
            'how': [
                'Select a preset scenario from the sidebar (Decline, Conservative, Base, Optimistic) or use the Scenario Lab for a custom configuration',
                'Toggle benchmark lines on/off using the checkboxes in the sidebar display options',
                'Use "Compare with scenarios" to overlay a second trajectory for direct visual comparison',
                'Hover over the chart to read exact GW values at any year',
            ],
            'answers': [
                'Where does global nuclear capacity head under stated government policies?',
                'How much does the choice of extension policy matter vs. new build rate?',
                'Which benchmark does our scenario most closely track, and from which year?',
            ],
        },
        {
            'num': '10.2', 'name': 'Capacity Breakdown',
            'what': (
                'Stacked area charts showing capacity decomposed by either region or technology type, '
                'year by year. The chart uses a bottom-up stack (individual unit contributions) '
                'bridging to a top-down stack post-2040. A "Region" view shows the 12 geographic '
                'groupings; a "Technology" view shows PWR, BWR, CANDU, VVER, RBMK, and other types.'
            ),
            'how': [
                'Switch between Region and Technology views using the toggle at the top of the chart',
                'Hover over the stacked areas — a unified tooltip shows all regions/technologies at that year',
                'The dashed line at 2040 marks the bottom-up / top-down transition',
                'Filter to specific regions using the Geography filter in the sidebar',
            ],
            'answers': [
                'Which regions drive growth or decline in each scenario?',
                'How does the technology mix change over time?',
                'When does the Asian fleet overtake the Western fleet in total capacity?',
            ],
        },
        {
            'num': '10.3', 'name': 'Retirements',
            'what': (
                'A stacked bar chart of annual capacity retirements (GW/year), broken down by '
                'region. Shows which years have the heaviest retirement burden and how the retirement '
                'profile changes across scenarios — AcceleratedRetirement front-loads retirements '
                'in the late 2020s and 2030s; ExtendedOperations pushes the same retirements into '
                'the 2040s and 2050s.'
            ),
            'how': [
                'Select different scenarios to compare retirement profiles',
                'Hover to identify peak retirement years and the regions most affected',
                'Use alongside the Global Projection tab to understand the "replacement" challenge',
            ],
            'answers': [
                'When are the biggest waves of retirements?',
                'How much new build is needed just to stand still (offset retirements)?',
                'Which regions face the steepest retirement cliffs?',
            ],
        },
        {
            'num': '10.4', 'name': 'New Additions',
            'what': (
                'Annual new capacity additions (GW/year) from the pipeline — under construction, '
                'planned, and proposed reactors entering service each year. Helps users see whether '
                'the pipeline is front-loaded (near-term heavy) or back-loaded.'
            ),
            'how': [
                'Adjust the pipeline preset (High/Medium/Low) in the Scenario Lab to see how additions change',
                'The construction delay adder shifts all bars to the right by the specified number of years',
                'Compare with the Retirements tab to see the net capacity balance year by year',
            ],
            'answers': [
                'When does the bulk of the Under Construction fleet come online?',
                'How much of the future growth depends on Proposed (speculative) projects?',
                'Which regions are leading new additions in the near term?',
            ],
        },
        {
            'num': '10.5', 'name': 'Country Snapshot',
            'what': (
                'A bar chart showing total nuclear capacity by country for a selected year, sorted '
                'from highest to lowest. Provides a quick read on which countries dominate the fleet '
                'at any point in the projection, and how the country ranking shifts across years '
                'and scenarios.'
            ),
            'how': [
                'Use the year selector at the top of the tab to choose the reference year (e.g., 2024, 2030, 2040, 2050)',
                'Switch scenarios in the sidebar to compare how each country\'s trajectory differs',
                'Apply the Geography filter to restrict the chart to a specific region',
                'Hover over bars to see exact GW values by country',
            ],
            'answers': [
                'Which countries have the most nuclear capacity at a given year?',
                'Has China surpassed the US in total nuclear capacity, and if so when?',
                'Which countries grow the most in absolute GW terms between 2024 and 2050?',
                'How does a specific country\'s ranking change across scenarios?',
            ],
        },
        {
            'num': '10.6', 'name': 'World Map',
            'what': (
                'A geographic map showing nuclear capacity by country at a selected year (default: 2024). '
                'Country color intensity represents capacity in GW — darker shading indicates higher '
                'capacity. The map provides a geographic perspective on the distribution of nuclear power '
                'and how it shifts across scenarios and years.'
            ),
            'how': [
                'Use the year slider to see how the geographic distribution evolves over time',
                'Switch scenarios in the sidebar to compare geographic patterns under different assumptions',
                'Hover over countries to read exact GW values',
            ],
            'answers': [
                'Which countries dominate nuclear capacity in 2030, 2040, 2050?',
                'Where are the biggest geographic shifts happening?',
                'How does China\'s fleet compare to the US fleet in each scenario?',
            ],
        },
        {
            'num': '10.7', 'name': 'Scenario Lab',
            'what': (
                'The full lever panel for building custom scenarios. Provides two complementary ways '
                'to configure a scenario — a plain-English Claude assistant and manual lever controls '
                '— that can be used together or independently. Allows independent control of extension '
                'policy, large reactor pipeline preset, SMR pipeline preset, construction delay adder, '
                'post-2040 new build rate, and unit-level reactor overrides.'
            ),
            'how_custom': True,  # signal to render custom how-to section
            'answers': [
                'What if France extends all its reactors to 60 years?',
                'What if the proposed Chinese pipeline is canceled entirely?',
                'How sensitive is the 2050 total to the post-2040 new build assumption?',
                'What capacity trajectory results from a user-defined combination of policies?',
            ],
        },
    ]

    for tab in tabs:
        heading(doc, f"{tab['num']}  {tab['name']}", level=2)
        body_para(doc, tab['what'])
        if tab.get('how_custom'):
            # Two-path Scenario Lab how-to
            heading(doc, 'How to use — two approaches (can be combined)', level=3)
            p_opt1 = doc.add_paragraph()
            r1 = p_opt1.add_run('Option 1: Claude Helper (plain English)  ')
            r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = TEAL
            r1b = p_opt1.add_run('— recommended for complex or multi-reactor changes')
            r1b.font.size = Pt(10); r1b.font.color.rgb = DARK_GREY
            bullet_para(doc, 'Click the Claude assistant panel inside the Scenario Lab tab')
            bullet_para(doc, 'Type your scenario intent in plain English, e.g.: "Retire all Canadian reactors by 2050, cancel all proposed European projects, and set post-2040 new build to 20 GW/yr"')
            bullet_para(doc, 'Claude generates the corresponding override entries and lever values automatically')
            bullet_para(doc, 'Review the suggested changes in the response, then click Accept to apply — the levers and What-If overrides update immediately across all tabs')
            bullet_para(doc, 'You can then fine-tune any individual value manually after accepting')
            bullet_para(doc, 'Best used for: reactor-level what-if overrides across many units, and exploring unfamiliar scenarios quickly — manual entry of individual reactor overrides for a large fleet can be time-consuming')
            p_opt2 = doc.add_paragraph()
            r2 = p_opt2.add_run('Option 2: Manual Lever & Override Controls')
            r2.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = BASE_C
            p_opt2.paragraph_format.space_before = Pt(8)
            bullet_para(doc, 'Use the macro levers (extension policy, pipeline preset, delay adder, post-2040 rate) for scenario-level changes that affect the whole fleet')
            bullet_para(doc, 'Use the What-If Builder panel to override individual reactor retirement years, capacity, or status — useful for specific known decisions (e.g., a single plant life extension approval)')
            bullet_para(doc, 'Add "synthetic builds" to insert hypothetical new reactors not in the database (by region, capacity per unit, units per year, start year, number of years)')
            bullet_para(doc, 'Use "Clear all selections" to reset all What-If overrides back to the base scenario without affecting the macro levers')
            bullet_para(doc, 'Click Apply Scenario to recompute — all other tabs update automatically')
        else:
            heading(doc, 'How to use', level=3)
            for item in tab['how']:
                bullet_para(doc, item)
        heading(doc, 'Key questions it answers', level=3)
        for item in tab['answers']:
            bullet_para(doc, item)
        doc.add_paragraph()

    # ── 10.8 DATA EXPORT & SCENARIO MANAGEMENT ────────────────────────────────
    heading(doc, '10.8  Data Export & Scenario Management', level=2)
    body_para(doc, (
        'The dashboard provides several ways to save, share, and export data — covering both '
        'model outputs and lever configurations.'
    ))
    export_rows = [
        ('Download Scenario + Settings',
         'Sidebar button (💾)',
         'Exports a 4-sheet Excel file: (1) lever settings for the current scenario, '
         '(2) global capacity projection year by year, (3) regional breakdown, '
         '(4) model description. Useful for sharing a full scenario package or archiving a configuration.'),
        ('Full Reactor List',
         'Header button (⬇)',
         'Downloads a complete Excel export of every reactor in the database — operational, '
         'pipeline, and long-term shutdown — with all metadata fields including retirement dates, '
         'tier, capacity, and expected online year.'),
        ('Per-chart data downloads',
         'Each chart tab',
         'Every chart includes a ⬇ Data button that downloads the underlying data table as CSV or Excel. '
         'Use this to extract the exact numbers behind any chart for offline analysis.'),
        ('Save / load custom scenario',
         'Scenario Lab',
         'Custom scenarios built in the Scenario Lab can be saved under a user-defined name and '
         'reloaded in a later session. Saved scenarios persist within the browser session; '
         'use the Download Scenario + Settings export to preserve a configuration across sessions.'),
    ]
    make_table(doc,
        ['Feature', 'Where to find it', 'What it contains'],
        export_rows,
        col_widths=[1.5, 1.2, 3.8],
        font_size=9,
    )

    page_break(doc)

    # ── APPENDIX A: GLOSSARY ──────────────────────────────────────────────────
    heading(doc, 'Appendix A: Glossary', level=1)
    glossary = [
        ('AcceleratedRetirement', 'Extension policy where no Tier 2 life extensions are granted; reactors retire at their baseline license expiry or declared date.'),
        ('Baseline Year', '2024 — the last full year for which IAEA PRIS provides a comprehensive global fleet snapshot.'),
        ('Bottom-up phase', 'The projection methodology used from 2024 to 2040, tracking every reactor unit individually.'),
        ('Construction delay adder', 'A scenario parameter (in years) added to all pipeline reactor expected online years to reflect systemic delays.'),
        ('CurrentPolicy', 'Extension policy where ~50% of eligible reactors receive one standard extension increment, reflecting historical approval rates.'),
        ('ExtendedOperations', 'Extension policy where all eligible reactors receive the maximum life extension permitted under current law.'),
        ('Expected online year', 'The projected year in which a pipeline reactor is expected to begin commercial operation, sourced from PRIS/WNA and adjusted by the construction delay adder.'),
        ('GW (gigawatt)', '1,000 megawatts of electrical generating capacity (net). Global nuclear fleet ~380 GW as of end-2024.'),
        ('IAEA PRIS', 'Power Reactor Information System — the IAEA\'s authoritative global reactor database, updated in real time and annually consolidated.'),
        ('IAEA RDS-1', 'Reference Data Series No. 1: Nuclear Power Reactors in the World — the annual IAEA benchmark publication providing Low/High scenario totals.'),
        ('IEA STEPS', 'Stated Policies Scenario — IEA\'s central scenario conditioned on enacted and announced (but not aspirational) government policies.'),
        ('IEA APS', 'Announced Pledges Scenario — IEA scenario conditioned on all national net-zero/energy pledges being met in full.'),
        ('is_smr', 'Database flag (Boolean) identifying reactors with net capacity ≤300 MW as Small Modular Reactors.'),
        ('Pipeline probability', 'A per-reactor score (0–1) reflecting project-specific confidence of completion. Used in custom scenarios; preset scenarios use binary 0/1 realization.'),
        ('Planned', 'Pipeline status for reactors with firm commitment and financing expected; typically 3–5 years from construction start.'),
        ('Post-2040 new build rate', 'A scenario parameter (GW/year globally) determining annual capacity additions in the top-down phase.'),
        ('Proposed', 'Pipeline status for early-stage announced reactors; significant permitting and financing uncertainty remains.'),
        ('Tier 1', 'Reactor retirement date from a publicly declared, legally binding, or regulatory-approved decision (taken as-is).'),
        ('Tier 2', 'Reactor retirement date modelled by applying country-specific license rules to the commercial operation date.'),
        ('Top-down phase', 'Post-2040 projection methodology using a global growth rate distributed regionally, calibrated to benchmarks.'),
        ('transition_year', 'The year at which the model switches from bottom-up to top-down methodology (2040 for all pre-defined scenarios).'),
        ('UC (Under Construction)', 'Pipeline status for reactors where major civil construction is underway.'),
        ('What-If Builder', 'The Scenario Lab sub-panel allowing unit-level overrides to individual reactor retirement years, capacity, and status.'),
        ('WNA', 'World Nuclear Association — maintains a reactor database supplementing IAEA PRIS, particularly for pipeline projects.'),
    ]
    make_table(doc,
        ['Term', 'Definition'],
        glossary,
        col_widths=[2.0, 4.5],
        font_size=9,
    )

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Nuclear Capacity Dashboard — Methodology & User Guide  |  Version 1.0  |  April 2026  |  Based on IAEA PRIS 2024 data vintage')
    r.font.size = Pt(8)
    r.font.color.rgb = MID_GREY
    r.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_PATH)
    print(f'Saved: {OUTPUT_PATH}')


if __name__ == '__main__':
    build_doc()
